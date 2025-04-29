import os
import re
import tempfile
from config import SUPPORTED_FILE_TYPES
import PyPDF2
import docx

def read_txts_from_folder(folder_path):
    """
    Read content from text files in the specified folder.
    
    Args:
        folder_path (str): Path to the folder containing text files
        
    Returns:
        list: List of dictionaries with file_name and content
    """
    file_data = []
    
    # Check if folder exists
    if not os.path.exists(folder_path):
        print(f"Warning: Folder path does not exist: {folder_path}")
        return file_data
    
    for file_name in os.listdir(folder_path):
        # Get file extension
        _, ext = os.path.splitext(file_name)
        ext = ext.lower().lstrip('.')
        
        if ext in SUPPORTED_FILE_TYPES:
            file_path = os.path.join(folder_path, file_name)
            try:
                content = extract_text_from_file(file_path, ext)
                if content:
                    file_data.append({"file_name": file_name, "content": content})
            except Exception as e:
                print(f"Error processing file {file_name}: {str(e)}")
    
    if not file_data:
        print(f"Warning: No supported files found in {folder_path}")
        
    return file_data

def extract_text_from_file(file_path, file_type=None):
    """
    Extract text content from a file based on its type
    
    Args:
        file_path (str): Path to the file
        file_type (str): Type/extension of the file (txt, pdf, docx)
        
    Returns:
        str: Extracted text content
    """
    if file_type is None:
        _, ext = os.path.splitext(file_path)
        file_type = ext.lower().lstrip('.')
    
    if file_type == 'txt':
        return read_text_file(file_path)
    elif file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def read_text_file(file_path):
    """Read text from a .txt file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Fallback to ISO-8859-1 encoding if UTF-8 fails
        with open(file_path, 'r', encoding='ISO-8859-1') as file:
            return file.read()

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file"""
    try:
        # We need to import PyPDF2 dynamically to avoid dependency issues if not installed
        text = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text() + "\n"
        
        return text
    except ImportError:
        print("PyPDF2 not installed. Install it using: pip install PyPDF2")
        return ""
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from a .docx file"""
    try:
        # Import docx dynamically to avoid dependency issues if not installed
        
        
        doc = docx.Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Also get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except ImportError:
        print("python-docx not installed. Install it using: pip install python-docx")
        return ""
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        return ""

def handle_uploaded_file(uploaded_file, upload_folder):
    """
    Handle an uploaded file from Streamlit
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        upload_folder (str): Folder to save the uploaded file
        
    Returns:
        tuple: (file_path, file_name, content)
    """
    # Create temp file path
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    file_name = uploaded_file.name
    
    # Save the file
    file_path = os.path.join(upload_folder, file_name)
    
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    # Extract text based on file type
    file_type = file_extension.lstrip('.')
    if file_type in SUPPORTED_FILE_TYPES:
        content = extract_text_from_file(file_path, file_type)
        return file_path, file_name, content
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

